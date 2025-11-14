from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

def add_horizontal_line(paragraph):
    """Add a horizontal line below a paragraph"""
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '12')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '2C5F8D')
    pBdr.append(bottom)
    pPr.append(pBdr)

def generate_professional_docx(resume_data, output_dir, filename):
    doc = Document()
    
    # Set margins
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)

    # Header - Name (Large, Bold, Dark Blue)
    info = resume_data.get('personal_info', {})
    name_para = doc.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_run = name_para.add_run(info.get('name', '').upper())
    name_run.font.size = Pt(24)
    name_run.font.bold = True
    name_run.font.color.rgb = RGBColor(44, 95, 141)  # Professional dark blue
    
    # Contact Info (Centered, Gray)
    contact_para = doc.add_paragraph()
    contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_run = contact_para.add_run(
        f"{info.get('email', '')} • {info.get('phone', '')} • {info.get('location', '')}"
    )
    contact_run.font.size = Pt(10)
    contact_run.font.color.rgb = RGBColor(80, 80, 80)
    
    doc.add_paragraph()  # Spacer

    # Professional Summary Section
    if resume_data.get('summary'):
        summary_heading = doc.add_paragraph()
        summary_run = summary_heading.add_run('PROFESSIONAL SUMMARY')
        summary_run.font.size = Pt(13)
        summary_run.font.bold = True
        summary_run.font.color.rgb = RGBColor(44, 95, 141)
        add_horizontal_line(summary_heading)
        
        summary_text = doc.add_paragraph(resume_data['summary'])
        summary_text.paragraph_format.space_after = Pt(12)
        for run in summary_text.runs:
            run.font.size = Pt(10)

    # Skills Section
    if resume_data.get('skills'):
        skills_heading = doc.add_paragraph()
        skills_run = skills_heading.add_run('CORE COMPETENCIES')
        skills_run.font.size = Pt(13)
        skills_run.font.bold = True
        skills_run.font.color.rgb = RGBColor(44, 95, 141)
        add_horizontal_line(skills_heading)
        
        skills_text = doc.add_paragraph(' • '.join(resume_data['skills']))
        skills_text.paragraph_format.space_after = Pt(12)
        for run in skills_text.runs:
            run.font.size = Pt(10)

    # Work Experience Section
    if resume_data.get('experience'):
        exp_heading = doc.add_paragraph()
        exp_run = exp_heading.add_run('PROFESSIONAL EXPERIENCE')
        exp_run.font.size = Pt(13)
        exp_run.font.bold = True
        exp_run.font.color.rgb = RGBColor(44, 95, 141)
        add_horizontal_line(exp_heading)
        
        for exp in resume_data['experience']:
            # Job Title (Bold, Dark Blue)
            job_para = doc.add_paragraph()
            job_run = job_para.add_run(exp.get('title', ''))
            job_run.font.size = Pt(11)
            job_run.font.bold = True
            job_run.font.color.rgb = RGBColor(44, 95, 141)
            
            # Company and Duration (Italic, Gray)
            company_para = doc.add_paragraph()
            company_run = company_para.add_run(
                f"{exp.get('company', '')} | {exp.get('duration', '')}"
            )
            company_run.font.size = Pt(10)
            company_run.font.italic = True
            company_run.font.color.rgb = RGBColor(100, 100, 100)
            
            # Description/Responsibilities
            if exp.get('description'):
                desc_para = doc.add_paragraph(exp['description'])
                desc_para.paragraph_format.left_indent = Inches(0.25)
                desc_para.paragraph_format.space_after = Pt(6)
                for run in desc_para.runs:
                    run.font.size = Pt(10)
            
            if exp.get('responsibilities'):
                for resp in exp['responsibilities']:
                    resp_para = doc.add_paragraph(f'• {resp}')
                    resp_para.paragraph_format.left_indent = Inches(0.25)
                    resp_para.paragraph_format.space_after = Pt(4)
                    for run in resp_para.runs:
                        run.font.size = Pt(10)
            
            doc.add_paragraph()  # Space between jobs

    # Education Section
    if resume_data.get('education'):
        edu_heading = doc.add_paragraph()
        edu_run = edu_heading.add_run('EDUCATION')
        edu_run.font.size = Pt(13)
        edu_run.font.bold = True
        edu_run.font.color.rgb = RGBColor(44, 95, 141)
        add_horizontal_line(edu_heading)
        
        for edu in resume_data['education']:
            degree_para = doc.add_paragraph()
            degree_run = degree_para.add_run(edu.get('degree', ''))
            degree_run.font.size = Pt(11)
            degree_run.font.bold = True
            
            institution_para = doc.add_paragraph()
            institution_run = institution_para.add_run(
                f"{edu.get('institution', '')} | {edu.get('year', '')}"
            )
            institution_run.font.size = Pt(10)
            institution_run.font.color.rgb = RGBColor(100, 100, 100)
            
            if edu.get('gpa'):
                gpa_para = doc.add_paragraph(f"GPA: {edu['gpa']}")
                gpa_para.paragraph_format.left_indent = Inches(0.25)
                for run in gpa_para.runs:
                    run.font.size = Pt(10)

    # Projects Section
    if resume_data.get('projects'):
        proj_heading = doc.add_paragraph()
        proj_run = proj_heading.add_run('PROJECTS')
        proj_run.font.size = Pt(13)
        proj_run.font.bold = True
        proj_run.font.color.rgb = RGBColor(44, 95, 141)
        add_horizontal_line(proj_heading)
        
        for proj in resume_data['projects']:
            proj_name_para = doc.add_paragraph()
            proj_name_run = proj_name_para.add_run(proj.get('name', ''))
            proj_name_run.font.size = Pt(11)
            proj_name_run.font.bold = True
            proj_name_run.font.color.rgb = RGBColor(44, 95, 141)
            
            proj_desc_para = doc.add_paragraph(proj.get('description', ''))
            proj_desc_para.paragraph_format.left_indent = Inches(0.25)
            proj_desc_para.paragraph_format.space_after = Pt(8)
            for run in proj_desc_para.runs:
                run.font.size = Pt(10)

    path = os.path.join(output_dir, filename)
    doc.save(path)
    return path