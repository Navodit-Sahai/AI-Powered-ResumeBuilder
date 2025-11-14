from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

def add_shaded_background(paragraph, color_rgb):
    """Add background shading to a paragraph"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color_rgb)
    paragraph._element.get_or_add_pPr().append(shading_elm)

def generate_modern_docx(resume_data, output_dir, filename):
    doc = Document()
    
    # Set narrow margins for modern look
    for section in doc.sections:
        section.top_margin = Inches(0.4)
        section.bottom_margin = Inches(0.4)
        section.left_margin = Inches(0.6)
        section.right_margin = Inches(0.6)

    # Header Section with colored background
    info = resume_data.get('personal_info', {})
    
    # Name (Extra Large, White on Dark Blue Background)
    name_para = doc.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_shaded_background(name_para, '1A4D7E')  # Dark blue background
    name_para.paragraph_format.space_before = Pt(6)
    name_para.paragraph_format.space_after = Pt(6)
    
    name_run = name_para.add_run(info.get('name', '').upper())
    name_run.font.size = Pt(28)
    name_run.font.bold = True
    name_run.font.color.rgb = RGBColor(255, 255, 255)  # White text
    
    # Contact Info (White on Blue)
    contact_para = doc.add_paragraph()
    contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_shaded_background(contact_para, '1A4D7E')
    contact_para.paragraph_format.space_after = Pt(8)
    
    contact_run = contact_para.add_run(
        f"📧 {info.get('email', '')}  •  📱 {info.get('phone', '')}  •  📍 {info.get('location', '')}"
    )
    contact_run.font.size = Pt(10)
    contact_run.font.color.rgb = RGBColor(255, 255, 255)
    
    doc.add_paragraph()  # Spacer

    # Professional Summary with accent box
    if resume_data.get('summary'):
        summary_heading = doc.add_paragraph()
        summary_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
        add_shaded_background(summary_heading, 'E8F4F8')  # Light blue background
        summary_heading.paragraph_format.space_before = Pt(4)
        summary_heading.paragraph_format.space_after = Pt(4)
        summary_heading.paragraph_format.left_indent = Inches(0.1)
        
        summary_run = summary_heading.add_run('💼 PROFESSIONAL PROFILE')
        summary_run.font.size = Pt(12)
        summary_run.font.bold = True
        summary_run.font.color.rgb = RGBColor(26, 77, 126)
        
        summary_text = doc.add_paragraph(resume_data['summary'])
        summary_text.paragraph_format.left_indent = Inches(0.15)
        summary_text.paragraph_format.space_after = Pt(12)
        for run in summary_text.runs:
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(60, 60, 60)

    # Skills Section with colored box
    if resume_data.get('skills'):
        skills_heading = doc.add_paragraph()
        add_shaded_background(skills_heading, 'E8F4F8')
        skills_heading.paragraph_format.space_before = Pt(4)
        skills_heading.paragraph_format.space_after = Pt(4)
        skills_heading.paragraph_format.left_indent = Inches(0.1)
        
        skills_run = skills_heading.add_run('⚡ CORE COMPETENCIES')
        skills_run.font.size = Pt(12)
        skills_run.font.bold = True
        skills_run.font.color.rgb = RGBColor(26, 77, 126)
        
        # Create skill pills
        skills_para = doc.add_paragraph()
        skills_para.paragraph_format.left_indent = Inches(0.15)
        skills_para.paragraph_format.space_after = Pt(12)
        
        for i, skill in enumerate(resume_data['skills']):
            skill_run = skills_para.add_run(f' {skill} ')
            skill_run.font.size = Pt(9)
            skill_run.font.bold = True
            skill_run.font.color.rgb = RGBColor(26, 77, 126)
            
            # Add separator
            if i < len(resume_data['skills']) - 1:
                sep_run = skills_para.add_run(' • ')
                sep_run.font.color.rgb = RGBColor(150, 150, 150)

    # Work Experience
    if resume_data.get('experience'):
        exp_heading = doc.add_paragraph()
        add_shaded_background(exp_heading, 'E8F4F8')
        exp_heading.paragraph_format.space_before = Pt(4)
        exp_heading.paragraph_format.space_after = Pt(4)
        exp_heading.paragraph_format.left_indent = Inches(0.1)
        
        exp_run = exp_heading.add_run('💼 PROFESSIONAL EXPERIENCE')
        exp_run.font.size = Pt(12)
        exp_run.font.bold = True
        exp_run.font.color.rgb = RGBColor(26, 77, 126)
        
        for exp in resume_data['experience']:
            # Job Title and Company in same line
            job_para = doc.add_paragraph()
            job_para.paragraph_format.left_indent = Inches(0.15)
            job_para.paragraph_format.space_before = Pt(8)
            
            title_run = job_para.add_run(exp.get('title', ''))
            title_run.font.size = Pt(11)
            title_run.font.bold = True
            title_run.font.color.rgb = RGBColor(26, 77, 126)
            
            job_para.add_run(' @ ')
            
            company_run = job_para.add_run(exp.get('company', ''))
            company_run.font.size = Pt(11)
            company_run.font.bold = True
            company_run.font.color.rgb = RGBColor(230, 90, 90)  # Red accent
            
            # Duration
            duration_para = doc.add_paragraph()
            duration_para.paragraph_format.left_indent = Inches(0.15)
            duration_run = duration_para.add_run(f"📅 {exp.get('duration', '')}")
            duration_run.font.size = Pt(9)
            duration_run.font.italic = True
            duration_run.font.color.rgb = RGBColor(120, 120, 120)
            
            # Description
            if exp.get('description'):
                desc_para = doc.add_paragraph(exp['description'])
                desc_para.paragraph_format.left_indent = Inches(0.3)
                desc_para.paragraph_format.space_after = Pt(4)
                for run in desc_para.runs:
                    run.font.size = Pt(10)
            
            # Responsibilities
            if exp.get('responsibilities'):
                for resp in exp['responsibilities']:
                    resp_para = doc.add_paragraph()
                    resp_para.paragraph_format.left_indent = Inches(0.3)
                    resp_para.paragraph_format.space_after = Pt(3)
                    
                    bullet_run = resp_para.add_run('▸ ')
                    bullet_run.font.color.rgb = RGBColor(26, 77, 126)
                    bullet_run.font.bold = True
                    
                    resp_run = resp_para.add_run(resp)
                    resp_run.font.size = Pt(10)

    # Education Section
    if resume_data.get('education'):
        edu_heading = doc.add_paragraph()
        add_shaded_background(edu_heading, 'E8F4F8')
        edu_heading.paragraph_format.space_before = Pt(4)
        edu_heading.paragraph_format.space_after = Pt(4)
        edu_heading.paragraph_format.left_indent = Inches(0.1)
        
        edu_run = edu_heading.add_run('🎓 EDUCATION')
        edu_run.font.size = Pt(12)
        edu_run.font.bold = True
        edu_run.font.color.rgb = RGBColor(26, 77, 126)
        
        for edu in resume_data['education']:
            degree_para = doc.add_paragraph()
            degree_para.paragraph_format.left_indent = Inches(0.15)
            degree_para.paragraph_format.space_before = Pt(6)
            
            degree_run = degree_para.add_run(edu.get('degree', ''))
            degree_run.font.size = Pt(11)
            degree_run.font.bold = True
            degree_run.font.color.rgb = RGBColor(26, 77, 126)
            
            institution_para = doc.add_paragraph()
            institution_para.paragraph_format.left_indent = Inches(0.15)
            institution_run = institution_para.add_run(
                f"{edu.get('institution', '')} • {edu.get('year', '')}"
            )
            institution_run.font.size = Pt(10)
            institution_run.font.color.rgb = RGBColor(100, 100, 100)

    # Projects Section
    if resume_data.get('projects'):
        proj_heading = doc.add_paragraph()
        add_shaded_background(proj_heading, 'E8F4F8')
        proj_heading.paragraph_format.space_before = Pt(4)
        proj_heading.paragraph_format.space_after = Pt(4)
        proj_heading.paragraph_format.left_indent = Inches(0.1)
        
        proj_run = proj_heading.add_run('🚀 NOTABLE PROJECTS')
        proj_run.font.size = Pt(12)
        proj_run.font.bold = True
        proj_run.font.color.rgb = RGBColor(26, 77, 126)
        
        for proj in resume_data['projects']:
            proj_name_para = doc.add_paragraph()
            proj_name_para.paragraph_format.left_indent = Inches(0.15)
            proj_name_para.paragraph_format.space_before = Pt(6)
            
            proj_name_run = proj_name_para.add_run(proj.get('name', ''))
            proj_name_run.font.size = Pt(11)
            proj_name_run.font.bold = True
            proj_name_run.font.color.rgb = RGBColor(26, 77, 126)
            
            proj_desc_para = doc.add_paragraph(proj.get('description', ''))
            proj_desc_para.paragraph_format.left_indent = Inches(0.3)
            proj_desc_para.paragraph_format.space_after = Pt(6)
            for run in proj_desc_para.runs:
                run.font.size = Pt(10)

    path = os.path.join(output_dir, filename)
    doc.save(path)
    return path