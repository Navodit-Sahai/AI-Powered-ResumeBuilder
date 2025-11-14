from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

def add_bottom_border(paragraph, color='2E4053', size='18'):
    """Add a bottom border to a paragraph"""
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), size)
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color)
    pBdr.append(bottom)
    pPr.append(pBdr)

def generate_academic_docx(resume_data, output_dir, filename):
    doc = Document()
    
    # Set traditional academic margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    info = resume_data.get('personal_info', {})
    
    # Name (Large, Centered, Dark Academic Color)
    name_para = doc.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_run = name_para.add_run(info.get('name', '').upper())
    name_run.font.size = Pt(22)
    name_run.font.bold = True
    name_run.font.color.rgb = RGBColor(46, 64, 83)  # Academic dark gray-blue
    add_bottom_border(name_para, '2E4053', '24')
    
    # Contact Information (Centered)
    contact_para = doc.add_paragraph()
    contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_para.paragraph_format.space_after = Pt(16)
    
    email_run = contact_para.add_run(info.get('email', ''))
    email_run.font.size = Pt(10)
    email_run.font.color.rgb = RGBColor(80, 80, 80)
    
    contact_para.add_run(' • ')
    
    phone_run = contact_para.add_run(info.get('phone', ''))
    phone_run.font.size = Pt(10)
    phone_run.font.color.rgb = RGBColor(80, 80, 80)
    
    if info.get('location'):
        contact_para.add_run(' • ')
        location_run = contact_para.add_run(info.get('location', ''))
        location_run.font.size = Pt(10)
        location_run.font.color.rgb = RGBColor(80, 80, 80)

    # Education Section (Most Important for Academic)
    if resume_data.get('education'):
        edu_heading = doc.add_paragraph('EDUCATION')
        edu_heading.paragraph_format.space_before = Pt(12)
        edu_heading.paragraph_format.space_after = Pt(8)
        edu_run = edu_heading.runs[0]
        edu_run.font.size = Pt(14)
        edu_run.font.bold = True
        edu_run.font.color.rgb = RGBColor(46, 64, 83)
        add_bottom_border(edu_heading, '8B4513', '12')  # Brown academic color
        
        for edu in resume_data['education']:
            # Degree (Bold)
            degree_para = doc.add_paragraph()
            degree_para.paragraph_format.space_before = Pt(8)
            degree_run = degree_para.add_run(edu.get('degree', ''))
            degree_run.font.size = Pt(11)
            degree_run.font.bold = True
            degree_run.font.color.rgb = RGBColor(46, 64, 83)
            
            # Institution and Year
            institution_para = doc.add_paragraph()
            institution_para.paragraph_format.left_indent = Inches(0.25)
            
            inst_run = institution_para.add_run(edu.get('institution', ''))
            inst_run.font.size = Pt(10)
            inst_run.font.italic = True
            
            institution_para.add_run(' • ')
            
            year_run = institution_para.add_run(edu.get('year', ''))
            year_run.font.size = Pt(10)
            year_run.font.italic = True
            year_run.font.color.rgb = RGBColor(100, 100, 100)
            
            # GPA if present
            if edu.get('gpa'):
                gpa_para = doc.add_paragraph()
                gpa_para.paragraph_format.left_indent = Inches(0.25)
                gpa_run = gpa_para.add_run(f"GPA: {edu['gpa']}")
                gpa_run.font.size = Pt(10)
                gpa_run.font.bold = True
                gpa_run.font.color.rgb = RGBColor(139, 69, 19)  # Academic brown

    # Research Experience / Work Experience
    if resume_data.get('experience'):
        exp_heading = doc.add_paragraph('RESEARCH & PROFESSIONAL EXPERIENCE')
        exp_heading.paragraph_format.space_before = Pt(16)
        exp_heading.paragraph_format.space_after = Pt(8)
        exp_run = exp_heading.runs[0]
        exp_run.font.size = Pt(14)
        exp_run.font.bold = True
        exp_run.font.color.rgb = RGBColor(46, 64, 83)
        add_bottom_border(exp_heading, '8B4513', '12')
        
        for exp in resume_data['experience']:
            # Position Title (Bold)
            title_para = doc.add_paragraph()
            title_para.paragraph_format.space_before = Pt(8)
            title_run = title_para.add_run(exp.get('title', ''))
            title_run.font.size = Pt(11)
            title_run.font.bold = True
            title_run.font.color.rgb = RGBColor(46, 64, 83)
            
            # Institution/Company and Duration
            company_para = doc.add_paragraph()
            company_para.paragraph_format.left_indent = Inches(0.25)
            
            company_run = company_para.add_run(exp.get('company', ''))
            company_run.font.size = Pt(10)
            company_run.font.italic = True
            
            company_para.add_run(' • ')
            
            duration_run = company_para.add_run(exp.get('duration', ''))
            duration_run.font.size = Pt(10)
            duration_run.font.italic = True
            duration_run.font.color.rgb = RGBColor(100, 100, 100)
            
            # Description
            if exp.get('description'):
                desc_para = doc.add_paragraph()
                desc_para.paragraph_format.left_indent = Inches(0.25)
                desc_para.paragraph_format.space_after = Pt(4)
                desc_run = desc_para.add_run(exp['description'])
                desc_run.font.size = Pt(10)
            
            # Responsibilities/Achievements
            if exp.get('responsibilities'):
                for resp in exp['responsibilities']:
                    resp_para = doc.add_paragraph()
                    resp_para.paragraph_format.left_indent = Inches(0.5)
                    resp_para.paragraph_format.first_line_indent = Inches(-0.15)
                    resp_para.paragraph_format.space_after = Pt(3)
                    
                    bullet_run = resp_para.add_run('◆ ')
                    bullet_run.font.color.rgb = RGBColor(139, 69, 19)
                    
                    resp_run = resp_para.add_run(resp)
                    resp_run.font.size = Pt(10)

    # Projects/Publications Section
    if resume_data.get('projects'):
        proj_heading = doc.add_paragraph('PROJECTS & RESEARCH')
        proj_heading.paragraph_format.space_before = Pt(16)
        proj_heading.paragraph_format.space_after = Pt(8)
        proj_run = proj_heading.runs[0]
        proj_run.font.size = Pt(14)
        proj_run.font.bold = True
        proj_run.font.color.rgb = RGBColor(46, 64, 83)
        add_bottom_border(proj_heading, '8B4513', '12')
        
        for proj in resume_data['projects']:
            # Project Name (Bold)
            proj_name_para = doc.add_paragraph()
            proj_name_para.paragraph_format.space_before = Pt(8)
            proj_name_run = proj_name_para.add_run(proj.get('name', ''))
            proj_name_run.font.size = Pt(11)
            proj_name_run.font.bold = True
            proj_name_run.font.color.rgb = RGBColor(46, 64, 83)
            
            # Project Description
            proj_desc_para = doc.add_paragraph()
            proj_desc_para.paragraph_format.left_indent = Inches(0.25)
            proj_desc_para.paragraph_format.space_after = Pt(6)
            proj_desc_run = proj_desc_para.add_run(proj.get('description', ''))
            proj_desc_run.font.size = Pt(10)

    # Skills & Technical Competencies
    if resume_data.get('skills'):
        skills_heading = doc.add_paragraph('TECHNICAL SKILLS')
        skills_heading.paragraph_format.space_before = Pt(16)
        skills_heading.paragraph_format.space_after = Pt(8)
        skills_run = skills_heading.runs[0]
        skills_run.font.size = Pt(14)
        skills_run.font.bold = True
        skills_run.font.color.rgb = RGBColor(46, 64, 83)
        add_bottom_border(skills_heading, '8B4513', '12')
        
        skills_para = doc.add_paragraph()
        skills_para.paragraph_format.left_indent = Inches(0.25)
        
        # Group skills with bullet points
        for i, skill in enumerate(resume_data['skills']):
            skill_run = skills_para.add_run(skill)
            skill_run.font.size = Pt(10)
            
            if i < len(resume_data['skills']) - 1:
                sep_run = skills_para.add_run(' • ')
                sep_run.font.color.rgb = RGBColor(139, 69, 19)

    # Professional Summary (if exists)
    if resume_data.get('summary'):
        summary_heading = doc.add_paragraph('PROFESSIONAL SUMMARY')
        summary_heading.paragraph_format.space_before = Pt(16)
        summary_heading.paragraph_format.space_after = Pt(8)
        summary_run = summary_heading.runs[0]
        summary_run.font.size = Pt(14)
        summary_run.font.bold = True
        summary_run.font.color.rgb = RGBColor(46, 64, 83)
        add_bottom_border(summary_heading, '8B4513', '12')
        
        summary_para = doc.add_paragraph(resume_data['summary'])
        summary_para.paragraph_format.left_indent = Inches(0.25)
        summary_para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        for run in summary_para.runs:
            run.font.size = Pt(10)

    path = os.path.join(output_dir, filename)
    doc.save(path)
    return path